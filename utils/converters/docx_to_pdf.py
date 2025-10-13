# import asyncio
# import os
# import pdfplumber
# import cloudinary
# import cloudinary.uploader
# import tempfile
# from docx import Document
# from reportlab.pdfgen import canvas
# from reportlab.lib.pagesizes import letter
# cloudinary.config(
#     cloud_name = os.getenv("CLOUDINARY_CLOUD_NAME"),
#     api_key=os.getenv("CLOUDINARY_API_KEY"),
#     api_secret=os.getenv("CLOUDINARY_API_SECRET")
# )


# async def convert_docx_to_pdf(file):
#     contents = await file.read()

    
#     def _convert():
#         with tempfile.NamedTemporaryFile(delete=False,suffix=".docx") as tmp:
#             tmp.write(contents)
#             tmp_path = tmp.name
            
#         doc = Document(tmp_path)
#         pdf_path = tmp_path.replace(".docx",".pdf")
        
#         c = canvas.Canvas(pdf_path,pagesize=letter)
#         width,height = letter
#         y = height - 50

#         for para in doc.paragraphs:
#             text = para.text
#             if text.strip():
#                 c.drawString(50, y, text)
#                 y -= 20
#                 if y < 50:
#                     c.showPage()
#                     y = height - 50

#         c.save()
        
#         return tmp_path,pdf_path
    
#     tmp_path,pdf_path = await asyncio.to_thread(_convert)
#     upload_result = cloudinary.uploader.upload(pdf_path, resource_type="raw")
#     os.remove(tmp_path)
#     os.remove(pdf_path)

#     return upload_result["secure_url"], os.path.basename(pdf_path)


import asyncio
import os
import pdfplumber
import cloudinary
import cloudinary.uploader
import tempfile
from docx import Document
from reportlab.pdfgen import canvas
from reportlab.lib.pagesizes import letter
from reportlab.lib.styles import getSampleStyleSheet
from reportlab.platypus import SimpleDocTemplate, Paragraph, Spacer
from reportlab.lib.units import inch

cloudinary.config(
    cloud_name = os.getenv("CLOUDINARY_CLOUD_NAME"),
    api_key=os.getenv("CLOUDINARY_API_KEY"),
    api_secret=os.getenv("CLOUDINARY_API_SECRET")
)


async def convert_docx_to_pdf(file):
    contents = await file.read()
    
    def _convert():
        with tempfile.NamedTemporaryFile(delete=False, suffix=".docx") as tmp:
            tmp.write(contents)
            tmp_path = tmp.name
        
        doc = Document(tmp_path)
        pdf_path = tmp_path.replace(".docx", ".pdf")
        
        # ✅ Use SimpleDocTemplate for better formatting
        pdf_doc = SimpleDocTemplate(pdf_path, pagesize=letter)
        styles = getSampleStyleSheet()
        story = []
        
        for para in doc.paragraphs:
            text = para.text.strip()
            if text:
                # ✅ Use Paragraph for proper text wrapping
                p = Paragraph(text, styles['Normal'])
                story.append(p)
                story.append(Spacer(1, 0.2 * inch))
        
        # ✅ Build the PDF
        pdf_doc.build(story)
        
        return tmp_path, pdf_path
    
    tmp_path, pdf_path = await asyncio.to_thread(_convert)
    
    # ✅ Wrap Cloudinary upload in asyncio.to_thread with error handling
    try:
        upload_result = await asyncio.to_thread(
            cloudinary.uploader.upload,
            pdf_path,
            resource_type="raw",
            # Optional: Add these parameters
            access_mode="public",  # Make sure file is publicly accessible
            format="pdf"
        )
    except Exception as e:
        print(f"Cloudinary upload failed: {str(e)}")
        os.remove(tmp_path)
        os.remove(pdf_path)
        raise
    
    os.remove(tmp_path)
    os.remove(pdf_path)

    return upload_result["secure_url"], os.path.basename(pdf_path)